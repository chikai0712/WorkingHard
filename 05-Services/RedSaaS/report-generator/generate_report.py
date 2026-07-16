#!/usr/bin/env python3
"""
RedSaaS Report Generator
從 DefectDojo 抓取漏洞資料，透過 Ollama 生成繁體中文報告草稿，輸出 Word 檔。

使用方式：
    python generate_report.py \
        --dd-url http://localhost:8001 \
        --dd-token YOUR_TOKEN \
        --engagement-id 1 \
        --output report.docx
"""

import argparse
import re
import sys
from datetime import datetime
from pathlib import Path

import requests
import ollama
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


SEVERITY_ORDER = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Info": 4}


def sanitize_text(text: str, max_len: int = 800) -> str:
    """移除非 ASCII / 非中日韓字元的雜訊，保留中文、英文、數字、標點。"""
    if not text:
        return ""
    # 保留：基本 ASCII 可見字元、CJK 統一漢字、全形標點、半形標點
    cleaned = re.sub(
        r"[^\x20-\x7E\u3000-\u303F\u4E00-\u9FFF\uFF00-\uFFEF\u2010-\u2027\u2030-\u205E\n]",
        "",
        text,
    )
    # 壓縮連續空白行
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    return cleaned[:max_len]


def sanitize_ai_output(text: str) -> str:
    """過濾 AI 輸出中混入的非中英文段落（泰文、緬甸文、阿拉伯文等）。"""
    if not text:
        return ""
    lines = text.splitlines()
    clean_lines = []
    for line in lines:
        # 若一行中超過 30% 是非 ASCII 且非 CJK 的字元，視為雜訊行
        non_cjk_ascii = re.findall(
            r"[^\x00-\x7F\u3000-\u303F\u4E00-\u9FFF\uFF00-\uFFEF]", line
        )
        if len(line) > 0 and len(non_cjk_ascii) / len(line) > 0.3:
            continue
        clean_lines.append(line)
    return "\n".join(clean_lines).strip()
SEVERITY_COLOR = {
    "Critical": RGBColor(0xC0, 0x00, 0x00),
    "High":     RGBColor(0xFF, 0x00, 0x00),
    "Medium":   RGBColor(0xFF, 0x99, 0x00),
    "Low":      RGBColor(0xFF, 0xCC, 0x00),
    "Info":     RGBColor(0x00, 0x70, 0xC0),
}


def fetch_findings(dd_url: str, token: str, engagement_id: int) -> list[dict]:
    """從 DefectDojo 抓取指定 engagement 的所有 findings。"""
    headers = {"Authorization": f"Token {token}"}
    findings = []
    offset = 0
    limit = 100

    while True:
        url = f"{dd_url}/api/v2/findings/?engagement={engagement_id}&limit={limit}&offset={offset}"
        resp = requests.get(url, headers=headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
        findings.extend(data["results"])
        if not data.get("next") or len(data["results"]) < limit:
            break
        offset += limit

    # 依 title + severity 去重，保留每種漏洞最新的一筆
    seen = {}
    for f in findings:
        key = (f.get("title", ""), f.get("severity", ""))
        if key not in seen:
            seen[key] = f
    findings = list(seen.values())

    return sorted(findings, key=lambda f: SEVERITY_ORDER.get(f.get("severity", "Info"), 99))


def fetch_engagement(dd_url: str, token: str, engagement_id: int) -> dict:
    """抓取 engagement 基本資訊。"""
    headers = {"Authorization": f"Token {token}"}
    resp = requests.get(f"{dd_url}/api/v2/engagements/{engagement_id}/", headers=headers, timeout=30)
    resp.raise_for_status()
    return resp.json()


def fetch_product(dd_url: str, token: str, product_id: int) -> dict:
    """抓取 product 基本資訊。"""
    headers = {"Authorization": f"Token {token}"}
    resp = requests.get(f"{dd_url}/api/v2/products/{product_id}/", headers=headers, timeout=30)
    resp.raise_for_status()
    return resp.json()


def translate_finding(finding: dict, model: str) -> dict:
    """用 Ollama 將單一 finding 的關鍵欄位轉成繁體中文描述段落。"""
    title = sanitize_text(finding.get("title", "未知漏洞"), 200)
    severity = finding.get("severity", "Info")
    description = sanitize_text(finding.get("description") or "", 800)
    mitigation = sanitize_text(finding.get("mitigation") or "", 600)
    impact = sanitize_text(finding.get("impact") or "", 400)
    url = (finding.get("url") or "").strip()
    steps_to_reproduce = sanitize_text(finding.get("steps_to_reproduce") or "", 600)

    steps_hint = steps_to_reproduce or "（請根據漏洞名稱與描述推斷重現方式）"
    prompt = f"""你是一位資深資安顧問，正在撰寫給台灣企業客戶的滲透測試報告。
重要規則：
- 全程只能使用繁體中文（Traditional Chinese）回答
- 不得使用英文以外的其他語言（禁止泰文、阿拉伯文、日文假名等）
- curl 指令、URL、技術名詞保持英文原文

請根據以下漏洞資訊，撰寫報告段落：

漏洞名稱：{title}
嚴重程度：{severity}
發現於：{url or "（未知）"}
原始描述：{description or "（無）"}
原始影響：{impact or "（無）"}
原始建議：{mitigation or "（無）"}
重現資訊：{steps_hint}

請直接輸出以下五個區塊（不要加其他說明文字）：

【漏洞描述】
（2-3 句，說明漏洞成因與技術細節）

【風險影響】
（1-2 句，說明攻擊者可造成的業務衝擊）

【重現步驟】
（條列 2-4 步，包含具體的 curl 或工具指令範例）

【修補建議】
（條列 2-3 點具體修補步驟）

【修復驗證】
（1-2 句，說明如何驗證修補完成）
"""

    response = ollama.chat(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        options={"temperature": 0.3},
    )
    msg = response.get("message") if isinstance(response, dict) else getattr(response, "message", None)
    if msg is None:
        raise ValueError(f"unexpected response structure: {response}")
    content = msg.get("content") if isinstance(msg, dict) else getattr(msg, "content", None)
    if content is None:
        raise ValueError(f"no content in message: {msg}")
    return {
        "title": title,
        "severity": severity,
        "url": url,
        "cve": finding.get("cve", ""),
        "steps_to_reproduce": steps_to_reproduce,
        "ai_content": sanitize_ai_output(content.strip()),
    }


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_severity_badge(doc: Document, severity: str):
    p = doc.add_paragraph()
    run = p.add_run(f"  {severity}  ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def build_report(
    doc: Document,
    product: dict,
    engagement: dict,
    findings_with_ai: list[dict],
    dd_url: str,
    engagement_id: int,
    target: str = "",
    tools: str = "Nmap, Nuclei v3.11",
):
    """組裝完整報告結構。"""
    product_name = product.get("name", "未知系統")
    eng_name = engagement.get("name", "資安評估")
    today = datetime.now().strftime("%Y 年 %m 月 %d 日")

    # 封面
    title_p = doc.add_heading(f"{product_name}", 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"{eng_name} — 滲透測試報告草稿")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    date_p = doc.add_paragraph(today)
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 測試範圍與環境
    add_heading(doc, "一、測試範圍與環境", 1)
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = "Table Grid"
    rows_data = [
        ("測試目標", target or engagement.get("target_url") or "—"),
        ("測試系統", product_name),
        ("Engagement 名稱", eng_name),
        ("測試期間", f"{engagement.get('target_start','—')} ～ {engagement.get('target_end','—')}"),
        ("使用工具", tools),
        ("測試方法論", "黑箱 / 灰箱滲透測試（OWASP API Security Top 10、OWASP Testing Guide）"),
    ]
    for i, (label, value) in enumerate(rows_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = value
    doc.add_paragraph("")
    doc.add_page_break()

    # 摘要統計
    add_heading(doc, "二、執行摘要", 1)
    counts = {}
    for f in findings_with_ai:
        counts[f["severity"]] = counts.get(f["severity"], 0) + 1

    summary_lines = [f"本次評估共發現 {len(findings_with_ai)} 個安全風險："]
    for sev in ["Critical", "High", "Medium", "Low", "Info"]:
        if sev in counts:
            summary_lines.append(f"  • {sev}（{sev_zh(sev)}）：{counts[sev]} 項")
    doc.add_paragraph("\n".join(summary_lines))

    doc.add_paragraph(
        "以下各章節依嚴重程度由高至低排列，每項漏洞包含技術描述、風險影響及修補建議。"
        "本報告為 AI 輔助生成的草稿，請資安顧問審閱後再交付客戶。"
    )
    doc.add_page_break()

    # 漏洞詳情
    add_heading(doc, "三、漏洞詳細說明", 1)

    if not findings_with_ai:
        doc.add_paragraph(
            "本次評估未發現需要回報的安全風險。掃描工具執行完畢，目標系統在本次測試範圍內"
            "未觸發任何已知漏洞特徵。建議定期重複執行評估以持續監控安全狀態。"
        )
    else:
        pass  # findings loop below

    for i, f in enumerate(findings_with_ai, 1):
        sev = f["severity"]
        add_heading(doc, f"{i}. {f['title']}", 2)

        meta = doc.add_paragraph()
        meta.add_run("嚴重程度：").bold = True
        color_run = meta.add_run(f"{sev}（{sev_zh(sev)}）")
        color_run.font.color.rgb = SEVERITY_COLOR.get(sev, RGBColor(0, 0, 0))
        color_run.bold = True

        if f.get("url"):
            url_p = doc.add_paragraph()
            url_p.add_run("影響端點：").bold = True
            url_p.add_run(f["url"])

        if f.get("cve"):
            cve_p = doc.add_paragraph()
            cve_p.add_run("CVE 編號：").bold = True
            cve_p.add_run(f["cve"])

        for block in f["ai_content"].split("\n\n"):
            if block.strip():
                p = doc.add_paragraph(block.strip())
                if block.strip().startswith("【") and p.runs:
                    p.runs[0].bold = True

        if f.get("steps_to_reproduce"):
            curl_heading = doc.add_paragraph()
            curl_heading.add_run("📋 重現指令（原始）").bold = True
            curl_p = doc.add_paragraph(f["steps_to_reproduce"])
            curl_p.paragraph_format.left_indent = Inches(0.3)
            for run in curl_p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)

        doc.add_paragraph("─" * 40)

    # 附錄
    add_heading(doc, "四、附錄", 1)
    doc.add_paragraph(
        f"DefectDojo 系統連結：{dd_url}/engagement/{engagement_id}/\n"
        "本報告由 RedSaaS Report Generator 自動生成，模型：llama3.2:3b\n"
        f"生成時間：{today}"
    )


def sev_zh(sev: str) -> str:
    mapping = {
        "Critical": "嚴重",
        "High": "高風險",
        "Medium": "中風險",
        "Low": "低風險",
        "Info": "資訊",
    }
    return mapping.get(sev, sev)


def main():
    parser = argparse.ArgumentParser(description="RedSaaS 中文滲透測試報告生成器")
    parser.add_argument("--dd-url", default="http://localhost:8001", help="DefectDojo URL")
    parser.add_argument("--dd-token", required=True, help="DefectDojo API Token")
    parser.add_argument("--engagement-id", type=int, required=True, help="Engagement ID")
    parser.add_argument("--output", default="report.docx", help="輸出檔案路徑")
    parser.add_argument("--model", default="llama3.2:3b", help="Ollama 模型名稱")
    parser.add_argument("--no-ai", action="store_true", help="跳過 AI 生成，只輸出原始資料")
    parser.add_argument("--target", default="", help="測試目標 URL 或 IP")
    parser.add_argument("--tools", default="Nmap, Nuclei v3.11", help="使用工具清單（逗號分隔）")
    args = parser.parse_args()

    print(f"[1/4] 連接 DefectDojo: {args.dd_url}")
    engagement = fetch_engagement(args.dd_url, args.dd_token, args.engagement_id)
    product = fetch_product(args.dd_url, args.dd_token, engagement["product"])
    print(f"      產品：{product['name']}，Engagement：{engagement['name']}")

    print(f"[2/4] 抓取 findings...")
    findings = fetch_findings(args.dd_url, args.dd_token, args.engagement_id)
    print(f"      共 {len(findings)} 個 findings")

    print(f"[3/4] AI 生成中文描述 (模型: {args.model})...")
    findings_with_ai = []
    for i, f in enumerate(findings, 1):
        title = f.get("title", "未知")
        print(f"      [{i}/{len(findings)}] {title[:50]}...")
        if args.no_ai:
            findings_with_ai.append({
                "title": title,
                "severity": f.get("severity", "Info"),
                "url": f.get("url", ""),
                "cve": f.get("cve", ""),
                "ai_content": f.get("description", "（無描述）")[:500],
            })
        else:
            try:
                result = translate_finding(f, args.model)
                findings_with_ai.append(result)
            except Exception as e:
                import traceback
                print(f"      [警告] AI 生成失敗: {e}")
                print(f"      [DEBUG] {traceback.format_exc()}")
                findings_with_ai.append({
                    "title": title,
                    "severity": f.get("severity", "Info"),
                    "url": f.get("url", ""),
                    "cve": f.get("cve", ""),
                    "ai_content": f.get("description", "（無描述）")[:500],
                })

    print(f"[4/4] 組裝 Word 報告...")
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    build_report(doc, product, engagement, findings_with_ai, args.dd_url, args.engagement_id,
                 target=args.target, tools=args.tools)

    output_path = Path(args.output)
    doc.save(output_path)
    print(f"\n完成！報告已儲存至：{output_path.absolute()}")


if __name__ == "__main__":
    main()
